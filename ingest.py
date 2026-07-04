import json
import re
from pathlib import Path


SOURCE_DIR = Path(
    r"C:\Users\admin\Desktop\хакатон норникель\Источники информации"
)
OUTPUT_DIR = Path("extracted")
INCLUDE_FOLDERS = ["Обзоры", "Статьи", "Доклады"]
MAX_FILE_SIZE = 20 * 1024 * 1024
PDF_MIN_TEXT_CHARS = 100


def clean_text(text):
    text = text.replace("\x00", "")
    text = re.sub(r"[ \t]+", " ", text)
    text = re.sub(r"\n{3,}", "\n\n", text)
    return text.strip()


def extract_pdf_text(path):
    try:
        import fitz
    except Exception as exc:
        raise RuntimeError("PyMuPDF is not installed or cannot be imported") from exc

    parts = []
    try:
        with fitz.open(path) as doc:
            for page_number, page in enumerate(doc, start=1):
                try:
                    parts.append(page.get_text("text"))
                except Exception as exc:
                    print(f"ERROR: failed to read PDF page {page_number}: {path} -> {exc}")
        return clean_text("\n".join(parts))
    except Exception:
        raise


def extract_docx_text(path):
    try:
        from docx import Document
    except Exception as exc:
        raise RuntimeError("python-docx is not installed or cannot be imported") from exc

    doc = Document(path)
    parts = []

    for paragraph in doc.paragraphs:
        text = paragraph.text.strip()
        if text:
            parts.append(text)

    for table in doc.tables:
        for row in table.rows:
            cells = [clean_text(cell.text) for cell in row.cells]
            row_text = " | ".join(cells).strip()
            if row_text:
                parts.append(row_text)

    return clean_text("\n".join(parts))


def relative_folder(path):
    try:
        return str(path.parent.relative_to(SOURCE_DIR))
    except Exception:
        return ""


def output_text_path(path):
    relative_path = path.relative_to(SOURCE_DIR)
    safe_name = "__".join(relative_path.parts) + ".txt"
    return OUTPUT_DIR / safe_name


def make_manifest_entry(path, status, text_size=0, reason=None):
    entry = {
        "file": path.name,
        "folder": relative_folder(path),
        "type": path.suffix.lower().lstrip("."),
        "text_size": text_size,
        "status": status,
    }
    if reason:
        entry["reason"] = reason
    return entry


def is_included_folder(path):
    relative_path = str(path.relative_to(SOURCE_DIR))
    return any(folder in relative_path for folder in INCLUDE_FOLDERS)


def save_text(path, text):
    target_path = output_text_path(path)
    try:
        target_path.parent.mkdir(parents=True, exist_ok=True)
        target_path.write_text(text, encoding="utf-8")
        print(f"  saved: {target_path}")
        return True
    except Exception as exc:
        print(f"ERROR: failed to save text for {path}: {exc}")
        return False


def process_file(path):
    print(f"Processing: {path}")
    suffix = path.suffix.lower()

    try:
        if not is_included_folder(path):
            print("  skipped: folder excluded")
            return make_manifest_entry(path, "skipped", reason="folder excluded")

        try:
            file_size = path.stat().st_size
        except Exception as exc:
            print(f"ERROR: failed to read file size: {path} -> {exc}")
            return make_manifest_entry(path, "error", reason=str(exc))

        if file_size > MAX_FILE_SIZE:
            print(f"  skipped: too large ({file_size} bytes)")
            return make_manifest_entry(path, "skipped", reason="too large")

        if suffix == ".pdf":
            text = extract_pdf_text(path)
            text_size = len(text)
            if text_size < PDF_MIN_TEXT_CHARS:
                print(f"  skipped: no text layer ({text_size} chars)")
                return make_manifest_entry(
                    path,
                    "skipped",
                    text_size=text_size,
                    reason="no text layer",
                )

            if not save_text(path, text):
                return make_manifest_entry(path, "error", text_size=text_size, reason="save failed")

            print(f"  ok: pdf, {text_size} chars")
            return make_manifest_entry(path, "ok", text_size=text_size)

        if suffix == ".docx":
            text = extract_docx_text(path)
            text_size = len(text)

            if not save_text(path, text):
                return make_manifest_entry(path, "error", text_size=text_size, reason="save failed")

            print(f"  ok: docx, {text_size} chars")
            return make_manifest_entry(path, "ok", text_size=text_size)

        print(f"  skipped: unsupported extension ({suffix or 'no extension'})")
        return make_manifest_entry(path, "skipped", reason="unsupported extension")

    except Exception as exc:
        print(f"ERROR: failed to process file: {path} -> {exc}")
        return make_manifest_entry(path, "error", reason=str(exc))


def write_manifest(manifest):
    manifest_path = OUTPUT_DIR / "manifest.json"
    try:
        OUTPUT_DIR.mkdir(parents=True, exist_ok=True)
        manifest_path.write_text(
            json.dumps(manifest, ensure_ascii=False, indent=2),
            encoding="utf-8",
        )
        print(f"Manifest saved: {manifest_path}")
    except Exception as exc:
        print(f"ERROR: failed to save manifest: {manifest_path} -> {exc}")


def print_summary(manifest):
    ok_count = sum(1 for item in manifest if item["status"] == "ok")
    skipped_count = sum(1 for item in manifest if item["status"] == "skipped")
    error_count = sum(1 for item in manifest if item["status"] == "error")

    print("")
    print("Summary:")
    print(f"  ok: {ok_count}")
    print(f"  skipped: {skipped_count}")
    print(f"  error: {error_count}")


def main():
    print(f"Source directory: {SOURCE_DIR}")
    print(f"Output directory: {OUTPUT_DIR}")
    print(f"Included folders: {', '.join(INCLUDE_FOLDERS)}")
    print("")

    if not SOURCE_DIR.exists():
        print(f"ERROR: source directory does not exist: {SOURCE_DIR}")
        return

    manifest = []

    try:
        paths = sorted(SOURCE_DIR.rglob("*"), key=lambda item: str(item).lower())
    except Exception as exc:
        print(f"ERROR: failed to scan source directory: {SOURCE_DIR} -> {exc}")
        return

    for path in paths:
        try:
            if not path.is_file():
                continue
        except Exception as exc:
            print(f"ERROR: failed to check file type: {path} -> {exc}")
            manifest.append(make_manifest_entry(path, "error", reason=str(exc)))
            continue

        manifest.append(process_file(path))

    write_manifest(manifest)
    print_summary(manifest)


if __name__ == "__main__":
    main()
